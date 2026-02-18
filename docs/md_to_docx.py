"""Script para convertir historia-usuario-agregar-empleado.md a .docx"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Título
doc.add_heading('Historia de Usuario: Agregar empleado con credenciales de acceso', 0)

# Descripción
p = doc.add_paragraph()
p.add_run('Como ').bold = False
p.add_run('administrador de RRHH').bold = True
p.add_run('\nQuiero ').bold = False
p.add_run('poder registrar nuevos empleados con sus datos personales y credenciales de acceso').bold = True
p.add_run('\nPara ').bold = False
p.add_run('que puedan ingresar al sistema y gestionar su información.').bold = True

doc.add_paragraph()

# Criterios de aceptación
doc.add_heading('Criterios de aceptación', level=1)

doc.add_paragraph('Dado que estoy logueado como administrador', style='List Bullet')
doc.add_paragraph('Cuando navego al módulo PIM > Add Employee', style='List Bullet')
doc.add_paragraph('Entonces debo ver el formulario de registro con los campos: First Name, Middle Name, Last Name, Employee Id, Create Login Details (switch), Username, Password, Confirm Password y Status.', style='List Bullet')

doc.add_paragraph()

doc.add_paragraph('Dado que estoy en el formulario Add Employee', style='List Bullet')
doc.add_paragraph('Cuando diligencio First Name, Middle Name, Last Name, Employee Id y activo Create Login Details con Username, Password y Status = Enabled', style='List Bullet')
doc.add_paragraph('Entonces al guardar debo ver un mensaje de éxito y el empleado debe aparecer en la lista de empleados.', style='List Bullet')

doc.add_paragraph()

doc.add_paragraph('Dado que un empleado fue registrado correctamente', style='List Bullet')
doc.add_paragraph('Cuando busco en Employee List por nombre (middle name) e Employee Id', style='List Bullet')
doc.add_paragraph('Entonces debo ver el empleado en la tabla de resultados con su Id, nombre y apellido.', style='List Bullet')

doc.add_paragraph()

# Definición de terminado
doc.add_heading('Definición de terminado', level=1)

doc.add_paragraph('☐ El empleado se guarda correctamente en la base de datos', style='List Bullet')
doc.add_paragraph('☐ Las credenciales permiten el login del nuevo empleado', style='List Bullet')
doc.add_paragraph('☐ El empleado aparece en la búsqueda de la lista de empleados', style='List Bullet')
doc.add_paragraph('☐ Los tests automatizados pasan', style='List Bullet')

doc.add_paragraph()

# Notas técnicas
doc.add_heading('Notas técnicas', level=1)

doc.add_paragraph('Prioridad: Alta', style='List Bullet')
doc.add_paragraph('Estimación: 5 puntos', style='List Bullet')
doc.add_paragraph('Feature: agregar_empleado.feature', style='List Bullet')

doc.save('docs/historia-usuario-agregar-empleado.docx')
print('Archivo generado: docs/historia-usuario-agregar-empleado.docx')
